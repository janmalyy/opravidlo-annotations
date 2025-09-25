import os
import json
import re

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from opravidlo_annotations.settings import FILES_DIR, DATA_CATEGORY, OPRAVIDLO_DIR


def print_json(data: dict, indent: int = 4) -> None:
    """
    Print JSON data with nice indentation.

    Args:
        data (dict): The JSON data to print.
        indent (int, optional): Number of spaces for indentation. Defaults to 4.
    """
    print(json.dumps(data, ensure_ascii=False, indent=indent))


def save_concordances_to_file(filename: str, concordances: list[str]) -> None:
    """
    Write concordances to a file. If the file does not exist, it will be created.
    Args:
        filename: filename to write to. Only the unique name, the prefix "data_zajmena" and the filename extension will be added.
        concordances: lines to be written to the file
    Returns: None
    """
    full_filename = FILES_DIR / f"{DATA_CATEGORY}_{filename}.txt"
    do_append = True
    if not full_filename.exists():
        do_append = False
        print(f"File {full_filename} does not exist, creating a new one.")

    with open(full_filename, "a" if do_append else "w", encoding="utf-8") as f:
        for c in concordances:
            f.write(c + "\n")

    print(f"Succesfully wrote {len(concordances)} concordances to {full_filename}.")


def set_document_language(document: docx.Document(), lang: str = "cs-CZ") -> None:
    """Set the default language for all text in the document.

    Args:
        document: A python-docx Document object.
        lang: Language code (e.g., "cs-CZ" for Czech).
    """
    styles = document.styles

    style = styles["Normal"]
    rpr = style.element.get_or_add_rPr()
    lang_elem = rpr.find(qn("w:lang"))

    if lang_elem is None:
        lang_elem = OxmlElement("w:lang")
        rpr.append(lang_elem)

    lang_elem.set(qn("w:val"), lang)

def save_concordances_to_word(concordances: list[str]) -> None:
    """
    Write concordances to the helper docx file.

    Args:
        concordances: lines to be written to the file

    Returns:
        None
    """
    doc = docx.Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Aptos"
    font.size = docx.shared.Pt(11)

    set_document_language(doc)

    for c in concordances:
        doc.add_paragraph(c)

    output_path = OPRAVIDLO_DIR / "opravidlo_annotations" / "files" / "word.docx"
    doc.save(output_path)
    os.startfile(output_path)

    print(f"Successfully wrote {len(concordances)} concordances to {output_path.name}.")


def count_correct_variants_in_json(filename: str) -> tuple[dict, int]:
    """
    Counts the occurrences of correct variants in a JSON file.

    Args:
        filename (str): Path to the JSON file.

    Returns:
        tuple[dict, int]: A dictionary with correct variants as keys and their total counts as values,
        and the total number of records.
    """
    filename = str(filename)
    if not filename.endswith(".json"):
        raise ValueError(f"File {filename} is not a JSON file.")

    with open(filename, "r", encoding="utf-8") as file:
        data = json.load(file)

    counter = {}

    for query in data.get("queries", []):
        correct_list = query.get("correct", [])
        if not correct_list:
            continue  # skip if 'correct' is empty or missing

        correct = correct_list[0]
        count = query.get("number_of_concordances", 0)
        counter[correct] = counter.get(correct, 0) + count

    total_number_of_records = sum(counter.values())
    return counter, total_number_of_records


def count_correct_variants_in_txt(filename: str) -> tuple[dict, int]:
    """
    Counts occurrences of bracketed variant expressions (e.g.[*vybyla|vybila|corpus*]),  in a text file.

    Args:
        filename (str): Path to the text file.

    Returns:
        tuple[dict, int]: A dictionary of bracketed expressions and their counts,
        and the total number of records.
    """
    filename = str(filename)
    if not filename.endswith(".txt"):
        raise ValueError(f"File {filename} is not a txt file.")

    counter = {}

    with open(filename, "r", encoding="utf-8") as file:
        for line in file:
            match = re.search(r'\[\*.*?\*\]', line)
            if match:
                key = match.group().lower()     # unify all to lowercase
                counter[key] = counter.get(key, 0) + 1

    total_number_of_records = sum(counter.values())
    return counter, total_number_of_records


def find_duplicates(strings: list[str]) -> tuple[list[str], list[str]]:
    """
    Find duplicates in a list of strings. The order remains unchanged.
    Args:
        strings (list[str]): List of strings to check.

    Returns:
        tuple[list[str], list[str]]: List of strings without duplicates and list of duplicates
        found in the input list; both without repetitions.
    """
    seen = []
    duplicates = []
    for string in strings:
        if string != "\n" and string in seen:   # we don't want to remove blank lines
            duplicates.append(string)
        else:
            seen.append(string)
    return seen, duplicates


def remove_duplicates(filename: str) -> None:
    """
    Returns: None; it saves the data back to the file without duplicates and prints the duplicates.
    """
    file_path = FILES_DIR / f"{DATA_CATEGORY}_{filename}.txt"
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    without_duplicates, duplicates = find_duplicates(lines)

    with open(file_path, "w", encoding="utf-8") as f:
        [f.write(line) for line in without_duplicates]

    if duplicates:
        print(f"Found {len(duplicates)} duplicates in '{filename}' file.")
        print([duplicate for duplicate in duplicates])
    else:
        print("No duplicates found.")


def check(filename: str) -> None:
    """
    Remove duplicates and then check the proportion of variants and whether the queries
    (and especially 'number of concordances' variable) are same in the JSON readme and in the real text.
    """
    remove_duplicates(filename)
    print("json: ", count_correct_variants_in_json(FILES_DIR / f"README_{filename}.json"))
    print("txt: ", count_correct_variants_in_txt(FILES_DIR / f"{DATA_CATEGORY}_{filename}.txt"))
    print()
